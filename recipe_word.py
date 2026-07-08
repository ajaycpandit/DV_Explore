"""recipe_word.py — additive bridge: build a formal, pharma/biotech-style DDS Word
document for a DeltaV BATCH_RECIPE / PROCEDURE FHX export.

The converter core (core/fhx_app.py) only ever builds Excel for the 'recipe' type — a
Word request silently returns the Excel workbook. This module fills that gap without
touching the byte-locked core, mirroring the visual style of core.build_dds_word
(cover page, revision history, numbered sections, sign-off table) but with a
recipe-appropriate structure: procedure narrative (SFC steps + transitions), formula
parameters, and formula value sets.

Data comes from recipe_bridge.parse_recipes() / parse_formulas(), which are already
used by the Recipes workspace, so the Word document stays consistent with the rest of
the app.
"""

import io
import datetime

import recipe_bridge


# ── low-level docx helpers (kept local so core stays untouched) ──────────────────
def _docx():
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    return Document, Pt, Cm, RGBColor, WD_ALIGN_PARAGRAPH, qn, OxmlElement


def _set_cell_bg(cell, hex_color, qn, OxmlElement):
    """Shade a table cell (same technique core uses via _set_cell_bg)."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


NAVY = '1F3864'
BLUE = '2E75B6'
GREY = '666666'
HDR_BG = '1F3864'
LBL_BG = 'D9E1F2'


def build_recipe_dds_word(text, fname, opts=None):
    """Return a BytesIO of a formal DDS-style Word document for the recipe(s) in `text`.

    opts (optional):
      include_formulas  (default True)  — include the formula value-set section
      include_procedure (default True)  — include the procedure narrative section
      max_steps         (default 0/all) — cap procedure rows for very large recipes
    """
    opts = opts or {}
    recs = recipe_bridge.parse_recipes(text)
    if not recs:
        raise ValueError('No BATCH_RECIPE / PROCEDURE found. Is this a Recipe FHX?')
    try:
        formulas = recipe_bridge.parse_formulas(text) or {}
    except Exception:
        formulas = {}

    Document, Pt, Cm, RGBColor, WD_ALIGN_PARAGRAPH, qn, OxmlElement = _docx()
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10)

    def heading(t, level=1, color=NAVY):
        p = doc.add_heading(t, level=level)
        if p.runs:
            p.runs[0].font.color.rgb = RGBColor.from_string(color)
            p.runs[0].font.name = 'Calibri'
        return p

    def para(t='', bold=False, italic=False, size=10, color=None):
        p = doc.add_paragraph()
        run = p.add_run(t)
        run.font.name = 'Calibri'
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
        return p

    def header_row(table, headers, bg=HDR_BG):
        hr = table.add_row()
        for i, h in enumerate(headers):
            hr.cells[i].text = str(h)
            _set_cell_bg(hr.cells[i], bg, qn, OxmlElement)
            for pr in hr.cells[i].paragraphs:
                for run in pr.runs:
                    run.font.bold = True
                    run.font.name = 'Calibri'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor.from_string('FFFFFF')
        return hr

    def data_row(table, cells, size=9, bold=False):
        row = table.add_row()
        for i, val in enumerate(cells):
            if i >= len(row.cells):
                break
            row.cells[i].text = '' if val is None else str(val)
            for pr in row.cells[i].paragraphs:
                for run in pr.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(size)
                    run.font.bold = bold
        return row

    # the primary recipe drives the cover / metadata
    primary = recs[0]
    meta = primary.get('meta', {})
    now = datetime.datetime.now()

    # ── COVER PAGE ───────────────────────────────────────────────────────────────
    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run('RECIPE DESIGN DESCRIPTION')
    r.font.name = 'Calibri'; r.font.size = Pt(20); r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub_p.add_run(meta.get('name', fname))
    r2.font.name = 'Calibri'; r2.font.size = Pt(14)
    r2.font.color.rgb = RGBColor.from_string(BLUE)

    type_p = doc.add_paragraph()
    type_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rtype = (meta.get('type') or 'PROCEDURE').replace('_', ' ').title()
    r3 = type_p.add_run('S88 ' + rtype)
    r3.font.name = 'Calibri'; r3.font.size = Pt(12); r3.font.italic = True
    r3.font.color.rgb = RGBColor.from_string(GREY)

    doc.add_paragraph()

    cover = doc.add_table(rows=0, cols=2)
    cover.style = 'Table Grid'
    cover_rows = [
        ('Document Number', (meta.get('name', fname) or fname) + '-RDD-001'),
        ('Recipe Name', meta.get('name', fname)),
        ('Recipe Type', rtype),
        ('Product', meta.get('product_name', '') or meta.get('product_code', '')),
        ('Recipe Version', meta.get('version', '')),
        ('Author', meta.get('author', '')),
        ('Document Type', 'Recipe Design Description'),
        ('Generated', now.strftime('%d-%b-%Y %H:%M')),
        ('Status', 'DRAFT'),
        ('Revision', '0.1'),
    ]
    for label, val in cover_rows:
        row = cover.add_row()
        row.cells[0].text = label
        row.cells[1].text = '' if val is None else str(val)
        for pr in row.cells[0].paragraphs:
            for run in pr.runs:
                run.font.bold = True; run.font.name = 'Calibri'; run.font.size = Pt(10)
        for pr in row.cells[1].paragraphs:
            for run in pr.runs:
                run.font.name = 'Calibri'; run.font.size = Pt(10)
        _set_cell_bg(row.cells[0], LBL_BG, qn, OxmlElement)
    cover.columns[0].width = Cm(6)
    cover.columns[1].width = Cm(10)

    doc.add_page_break()

    # ── 1. PURPOSE & SCOPE ───────────────────────────────────────────────────────
    heading('1. Purpose and Scope', 1)
    para('This document describes the design of the DeltaV batch recipe identified above, '
         'as extracted from its FHX export. It captures the procedural logic, formula '
         'parameters, and formula value sets that define the recipe, to support review, '
         'verification, and change control in a GxP context.')
    heading('1.1 Recipe Abstract', 2)
    abstract = (meta.get('abstract') or '').strip()
    para(abstract if abstract else 'No abstract was recorded in the recipe export.',
         italic=not abstract)
    if len(recs) > 1:
        para(f'This export contains {len(recs)} recipe objects. The procedural and formula '
             f'detail below covers each in turn.', size=9, color=GREY)

    # ── 2. REVISION HISTORY ──────────────────────────────────────────────────────
    heading('2. Revision History', 1)
    rev = doc.add_table(rows=0, cols=4)
    rev.style = 'Table Grid'
    header_row(rev, ['Rev', 'Date', 'Author', 'Description of Change'])
    data_row(rev, ['0.1', now.strftime('%d-%b-%Y'), 'Auto-generated',
                   'Initial draft generated from DeltaV recipe FHX export'], size=10)
    # if the recipe carries its own version, note it as context
    if meta.get('version'):
        data_row(rev, [meta.get('version'), '', meta.get('author', ''),
                       'Recipe version recorded in DeltaV at time of export'], size=10)

    doc.add_page_break()

    # ── per-recipe sections ──────────────────────────────────────────────────────
    section_no = 3
    for ri, rec in enumerate(recs):
        rmeta = rec.get('meta', {})
        rname = rmeta.get('name', fname)
        if len(recs) > 1:
            heading(f'{section_no}. Recipe: {rname}', 1)
        else:
            heading(f'{section_no}. Recipe Overview', 1)

        # overview facts
        ov = doc.add_table(rows=0, cols=4)
        ov.style = 'Table Grid'
        facts = [
            ('Name', rmeta.get('name', ''), 'Type', (rmeta.get('type') or '').title()),
            ('Product', rmeta.get('product_name', ''), 'Product Code', rmeta.get('product_code', '')),
            ('Version', rmeta.get('version', ''), 'Author', rmeta.get('author', '')),
            ('Default Batch', rmeta.get('default_batch_size', ''),
             'Batch Range', _batch_range(rmeta)),
        ]
        for a, b, c, d in facts:
            row = ov.add_row()
            vals = [a, b, c, d]
            for i, v in enumerate(vals):
                row.cells[i].text = '' if v is None else str(v)
                for pr in row.cells[i].paragraphs:
                    for run in pr.runs:
                        run.font.name = 'Calibri'; run.font.size = Pt(9)
                        run.font.bold = (i % 2 == 0)
            _set_cell_bg(row.cells[0], LBL_BG, qn, OxmlElement)
            _set_cell_bg(row.cells[2], LBL_BG, qn, OxmlElement)

        params = rec.get('params', []) or []
        proc = rec.get('procedure', {}) or {}
        steps = proc.get('steps', {}) or {}
        trans = proc.get('transitions', {}) or {}
        counts = para(
            f'This recipe defines {len(params)} formula parameter(s), '
            f'{len(steps)} procedural step(s), and {len(trans)} transition(s).',
            size=9, color=GREY)

        # ── procedure narrative ──
        if opts.get('include_procedure', True) and steps:
            heading(f'{section_no}.1  Procedural Logic', 2)
            para('The table below lists the recipe\u2019s procedural steps in sequence with '
                 'their outgoing transitions, describing the S88 procedure flow.', size=9)
            _procedure_table(doc, proc, opts, header_row, data_row, Cm)

        # ── parameters ──
        if params:
            heading(f'{section_no}.2  Formula Parameters', 2)
            _params_table(doc, params, header_row, data_row, Cm)

        # ── formulas ──
        if opts.get('include_formulas', True):
            fsets = formulas.get(rname) or formulas.get(rmeta.get('name', '')) or []
            if fsets:
                heading(f'{section_no}.3  Formula Value Sets', 2)
                para(f'{len(fsets)} formula value set(s) are defined for this recipe. Each '
                     'column below is one released formula; rows are parameters that carry a '
                     'configured value.', size=9)
                _formula_table(doc, params, fsets, header_row, data_row, Cm)

        section_no += 1
        if ri < len(recs) - 1:
            doc.add_page_break()

    # ── FINAL: REVIEW & APPROVAL ─────────────────────────────────────────────────
    doc.add_page_break()
    heading(f'{section_no}. Review and Approval', 1)
    para('This document requires review and approval before use in a GxP context.')
    doc.add_paragraph()
    sig = doc.add_table(rows=0, cols=4)
    sig.style = 'Table Grid'
    header_row(sig, ['Role', 'Name', 'Signature', 'Date'])
    for role in ['Author', 'Technical Reviewer', 'Quality Reviewer', 'Approver']:
        row = sig.add_row()
        row.cells[0].text = role
        for pr in row.cells[0].paragraphs:
            for run in pr.runs:
                run.font.name = 'Calibri'; run.font.size = Pt(10)
        for c in (1, 2, 3):
            for pr in row.cells[c].paragraphs:
                pr.add_run('  ')

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── section builders ─────────────────────────────────────────────────────────────
def _batch_range(meta):
    lo = meta.get('min_batch_size', '')
    hi = meta.get('max_batch_size', '')
    if lo or hi:
        return f'{lo} \u2013 {hi}'.strip(' \u2013')
    return ''


def _procedure_table(doc, proc, opts, header_row, data_row, Cm):
    steps = proc.get('steps', {}) or {}
    s2t = proc.get('s2t', []) or []
    # map step -> list of outgoing transitions
    out = {}
    for s, t in s2t:
        out.setdefault(s, []).append(t)
    tbl = doc.add_table(rows=0, cols=4)
    tbl.style = 'Table Grid'
    header_row(tbl, ['Seq', 'Step', 'Description', 'Outgoing Transition(s)'])
    max_steps = opts.get('max_steps', 0) or 0
    seq = 0
    for name, sdef in steps.items():
        seq += 1
        if max_steps and seq > max_steps:
            data_row(tbl, ['\u2026', f'(+{len(steps) - max_steps} more steps)',
                           'Truncated for length \u2014 see the Excel export for the full list', ''])
            break
        desc = (sdef.get('desc') or '').strip()
        outs = out.get(name, [])
        data_row(tbl, [seq, name, desc, ', '.join(outs)])
    for w, col in zip((1.6, 5.0, 6.0, 4.0), tbl.columns):
        col.width = Cm(w)


def _params_table(doc, params, header_row, data_row, Cm):
    tbl = doc.add_table(rows=0, cols=6)
    tbl.style = 'Table Grid'
    header_row(tbl, ['Parameter', 'Type', 'Default', 'Units', 'Range', 'Description'])
    for p in params:
        rng = ''
        lo, hi = p.get('low', ''), p.get('high', '')
        if lo != '' or hi != '':
            rng = f'{lo} \u2013 {hi}'.strip(' \u2013')
        data_row(tbl, [
            p.get('name', ''),
            p.get('param_type', '') or p.get('type', ''),
            p.get('value', '') if p.get('value', '') != '' else p.get('default', ''),
            p.get('units', ''),
            rng,
            (p.get('description', '') or '')[:120],
        ])
    for w, col in zip((3.8, 2.2, 2.2, 1.6, 2.2, 4.0), tbl.columns):
        col.width = Cm(w)


def _formula_table(doc, params, fsets, header_row, data_row, Cm):
    # limit columns to keep the table page-friendly; note overflow
    max_cols = 4
    shown = fsets[:max_cols]
    headers = ['Parameter'] + [fs.get('name', f'Set {i + 1}') for i, fs in enumerate(shown)]
    tbl = doc.add_table(rows=0, cols=len(headers))
    tbl.style = 'Table Grid'
    header_row(tbl, headers)
    # union of parameter names that appear in any shown set, preserving recipe param order
    ordered = [p.get('name', '') for p in params if p.get('name')]
    seen = set(ordered)
    for fs in shown:
        for k in (fs.get('values') or {}):
            if k not in seen:
                ordered.append(k); seen.add(k)
    for pname in ordered:
        cells = [pname]
        any_val = False
        for fs in shown:
            v = (fs.get('values') or {}).get(pname, '')
            if v != '':
                any_val = True
            cells.append(v)
        if any_val:
            data_row(tbl, cells)
    if len(fsets) > max_cols:
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run(f'Note: {len(fsets) - max_cols} additional formula value set(s) are '
                        f'defined but omitted here for page width. The full set of formulas is '
                        f'available in the Excel export (FORMULA VALUES sheet).')
        run.font.name = 'Calibri'; run.font.size = None
